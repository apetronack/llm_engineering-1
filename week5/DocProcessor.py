import pandas as pd
import numpy as np
from docx import Document
import os
from typing import Dict, List, Tuple
import logging

class DocProcessor:
    def __init__(self, max_table_preview_rows: int = 5):
        self.text_chunks = []
        self.tables = []
        self.logger = logging.getLogger(__name__)
        self.max_table_preview_rows = max_table_preview_rows
        
    def process_document(self, file_path: str) -> Tuple[List[str], List[pd.DataFrame]]:
        """
        Process a Word document, extracting both text and tables.
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            Tuple containing:
            - List of text chunks
            - List of pandas DataFrames (one per table)
        """
        doc = Document(file_path)
        
        # Process each paragraph and table in order of appearance
        for element in doc.element.body:
            if element.tag.endswith('p'):
                # Handle paragraphs
                paragraph = element.text.strip()
                if paragraph and len(paragraph) >= 50:  # Filter out very short paragraphs
                    self.text_chunks.append(paragraph)
                    
            elif element.tag.endswith('tbl'):
                try:
                    # Handle tables
                    table_data = []
                    max_cols = 0
                    
                    # First pass: collect all cells and determine max columns
                    for row in element.findall('.//w:tr', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                        row_data = []
                        for cell in row.findall('.//w:t', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                            row_data.append(cell.text.strip())
                        max_cols = max(max_cols, len(row_data))
                        table_data.append(row_data)
                    
                    if not table_data or max_cols == 0:
                        continue
                        
                    # Second pass: pad rows with empty strings if needed
                    padded_data = []
                    for row in table_data:
                        if len(row) < max_cols:
                            padded_row = row + [''] * (max_cols - len(row))
                            padded_data.append(padded_row)
                        else:
                            padded_data.append(row)
                    
                    # Generate column names if first row is empty or has fewer columns
                    if not padded_data[0] or any(not col for col in padded_data[0]):
                        columns = [f'Column_{i+1}' for i in range(max_cols)]
                        df = pd.DataFrame(padded_data, columns=columns)
                    else:
                        # Use first row as headers, replacing empty headers
                        headers = []
                        for i, header in enumerate(padded_data[0]):
                            if not header:
                                header = f'Column_{i+1}'
                            while header in headers:
                                header = f'{header}_dup'
                            headers.append(header)
                        
                        df = pd.DataFrame(padded_data[1:], columns=headers)
                    
                    # Clean empty cells using updated pandas methods
                    # Replace whitespace-only cells with None
                    df = df.apply(lambda x: x.str.strip() if x.dtype == "object" else x)
                    df = df.apply(lambda x: x.where(x.str.len() > 0, None) if x.dtype == "object" else x)
                    
                    # Drop rows where all values are None
                    df = df.dropna(how='all')
                    
                    # Convert None back to empty strings for consistent handling
                    df = df.fillna('')
                    
                    if not df.empty:
                        self.tables.append(df)
                        
                        # Create condensed table context for text embedding
                        preview_rows = min(len(df), self.max_table_preview_rows)
                        table_context = (
                            f"Table summary: {len(df)} rows with columns: {', '.join(df.columns)}. "
                            f"First {preview_rows} rows preview: \n"
                            f"{df.head(preview_rows).to_string(index=False)}"
                        )
                        self.text_chunks.append(table_context)
                        
                except Exception as e:
                    self.logger.warning(f"Error processing table in {file_path}: {str(e)}")
                    continue
        
        return self.text_chunks, self.tables

    def get_document_metadata(self, file_path: str) -> Dict:
        """
        Extract essential metadata from the Word document.
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            Dictionary containing essential document metadata
        """
        doc = Document(file_path)
        core_properties = doc.core_properties
        
        # Get filename without extension for fallback title
        filename = os.path.splitext(os.path.basename(file_path))[0]
        
        # Extract first paragraph for potential title if core title is empty
        first_paragraph = doc.paragraphs[0].text.strip() if doc.paragraphs else ""
        
        # Only include essential metadata fields
        metadata = {
            'title': core_properties.title or first_paragraph or filename,
            'author': core_properties.author or "Unknown",
            'created': core_properties.created,
            'modified': core_properties.modified,
            'filename': filename,
            'file_extension': os.path.splitext(file_path)[1],
            'num_tables': len(doc.tables),
            'category': core_properties.category or "Uncategorized",
            'last_modified_by': core_properties.last_modified_by or "Unknown"
        }
        
        # Clean up any None values
        return {k: str(v) if v is not None else "Not specified" for k, v in metadata.items()}

    def clear(self):
        """Reset the processor state"""
        self.text_chunks = []
        self.tables = []